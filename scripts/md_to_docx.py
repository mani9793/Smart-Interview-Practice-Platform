"""
One-off script: convert docs/TASK_SPLIT_PHASE2.md to docs/TASK_SPLIT_PHASE2.docx.
Run from project root: python scripts/md_to_docx.py
Requires: pip install python-docx
"""
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Run: pip install python-docx")
    raise

PROJECT_ROOT = Path(__file__).resolve().parent.parent
MD_FILE = PROJECT_ROOT / "docs" / "TASK_SPLIT_PHASE2.md"
DOCX_FILE = PROJECT_ROOT / "docs" / "TASK_SPLIT_PHASE2.docx"


def strip_bold(text):
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text)


def add_paragraph(doc, text, style=None, bold_runs=None):
    p = doc.add_paragraph(style=style)
    if not bold_runs:
        p.add_run(text)
        return p
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2] + ' ')
            run.bold = True
        else:
            p.add_run(part)
    return p


def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line or not line.startswith('|'):
            continue
        cells = [c.strip().strip('|') for c in line.split('|') if c.strip() != '']
        if not cells:
            continue
        if all(re.match(r'^[-:\s]+$', c) for c in cells):
            continue
        rows.append(cells)
    return rows


def main():
    doc = Document()
    doc.add_heading('Phase 2 Task Split — Two People', 0)

    content = MD_FILE.read_text(encoding='utf-8')
    lines = content.splitlines()

    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                rows = parse_table(table_lines)
                if rows:
                    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    t.style = 'Table Grid'
                    for ri, row_cells in enumerate(rows):
                        for ci, cell_text in enumerate(row_cells):
                            if ci < len(t.rows[ri].cells):
                                t.rows[ri].cells[ci].text = strip_bold(cell_text)
                    doc.add_paragraph()
            table_lines = []

        if line.strip() == '---':
            doc.add_paragraph('_' * 40)
            i += 1
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.strip().startswith('- '):
            text = line.strip()[2:]
            add_paragraph(doc, text, style='List Bullet')
            i += 1
            continue

        m = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if m:
            add_paragraph(doc, m.group(2), style='List Number')
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        add_paragraph(doc, line.strip())
        i += 1

    if in_table and table_lines:
        rows = parse_table(table_lines)
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = 'Table Grid'
            for ri, row_cells in enumerate(rows):
                for ci, cell_text in enumerate(row_cells):
                    if ci < len(t.rows[ri].cells):
                        t.rows[ri].cells[ci].text = strip_bold(cell_text)
            doc.add_paragraph()

    DOCX_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_FILE))
    print(f"Saved: {DOCX_FILE}")


if __name__ == '__main__':
    main()
